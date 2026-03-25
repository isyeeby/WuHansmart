from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建新文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题
title = doc.add_heading('Airbnb 房源价格预测：人工智能与传统方法的比较分析', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 作者信息
author = doc.add_paragraph('Nicola Camatti, Giacomo di Tollo, Gianni Filograsso, Sara Ghilardi')
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 期刊信息
journal = doc.add_paragraph('Computational Management Science (2024) 21:30')
journal.alignment = WD_ALIGN_PARAGRAPH.CENTER
journal.runs[0].italic = True

# DOI
doi = doc.add_paragraph('https://doi.org/10.1007/s10287-024-00511-4')
doi.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 摘要部分
doc.add_heading('摘要', level=2)
abstract = doc.add_paragraph(
    'Airbnb 的共享商业模式影响了广泛的利益群体，包括专业房东和传统住宿经营者，形成了一个复杂的生态系统。'
    '预测工作因此面临挑战。本文探讨 Airbnb 房源定价问题，提出一种计算方法与传统回归分析进行比较。'
    '研究使用意大利实际数据，聚焦于 2019 年 9 月的活跃 Airbnb 房源。通过收集不同房源产品的详细历史数据，'
    '分析各因素对价格的影响。研究采用多种机器学习算法，提供超越传统方法的全面视角。'
    '研究结果显示人工智能方法在预测准确性方面表现优异，但传统方法同样具有价值，特别是在可解释性方面。'
    '综合使用这些方法可为优化 Airbnb 市场策略提供有价值的参考。'
)

# 关键词
keywords = doc.add_paragraph()
keywords.add_run('关键词：').bold = True
keywords.add_run('Airbnb、机器学习、计算定价模型、价格预测系统、酒店业深度学习')

# 保存文档
output_path = 'D:\\code\\TujiaFeature\\docs\\Airbnb 房源分析_已排版.docx'
doc.save(output_path)
print('Document created:', output_path)
